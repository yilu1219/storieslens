from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "STORIESLENS_HANDOFF.md"
OUTPUT = ROOT / "exports" / "StoriesLens_Development_Handoff.docx"

NAVY = RGBColor(11, 29, 62)
BLUE = RGBColor(45, 98, 221)
PURPLE = RGBColor(123, 92, 238)
INK = RGBColor(30, 41, 59)
MUTED = RGBColor(91, 105, 126)
PALE = "EEF3FF"
LIGHT = "F6F8FC"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def format_inline(paragraph, text, color=INK, size=10.25, bold=False):
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        is_bold = part.startswith("**") and part.endswith("**")
        value = part[1:-1] if is_code else part[2:-2] if is_bold else part
        run = paragraph.add_run(value)
        run.font.name = "Consolas" if is_code else "Aptos"
        run.font.size = Pt(9.25 if is_code else size)
        run.font.color.rgb = NAVY if is_code else color
        run.bold = bold or is_bold
        if is_code:
            run.font.highlight_color = None


def add_callout(doc, text, label="Important"):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.name = "Aptos"
    run.font.color.rgb = BLUE
    format_inline(p, text, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 15, 6),
        ("Heading 2", 13.5, BLUE, 11, 4),
        ("Heading 3", 11.5, PURPLE, 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.25)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "STORIESLENS  |  DEVELOPMENT HANDOFF"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = MUTED
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("STORIESLENS")
    run.font.name = "Aptos Display"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = PURPLE

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Development Handoff")
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    format_inline(
        subtitle,
        "Product architecture, local setup, APIs, deployment, data migration, and next priorities",
        color=MUTED,
        size=12,
    )

    table = doc.add_table(rows=4, cols=2)
    set_table_fixed_width(table, [2200, 7160])
    table.style = "Table Grid"
    details = [
        ("Repository", "https://github.com/yilu1219/storieslens"),
        ("Production", "https://www.storieslens.com"),
        ("Primary branch", "main"),
        ("Updated", "2026-07-23"),
    ]
    for index, (label, value) in enumerate(details):
        left, right = table.rows[index].cells
        set_cell_shading(left, PALE)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        left.paragraphs[0].text = label
        right.paragraphs[0].text = value
        for run in left.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Aptos"
            run.font.color.rgb = NAVY
        for run in right.paragraphs[0].runs:
            run.font.name = "Aptos"
            run.font.color.rgb = INK

    doc.add_paragraph()
    add_callout(
        doc,
        "Rotate the OpenRouter key that appeared in development conversations. Store the replacement only in local and Railway environment variables.",
        "Security action",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    format_inline(
        p,
        "Purpose: let a new maintainer restore the project on another computer without relying on chat history.",
        color=MUTED,
        size=10,
    )
    doc.add_page_break()


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        parsed.append([cell.strip() for cell in row.strip().strip("|").split("|")])
    if len(parsed) < 2:
        return
    header = parsed[0]
    body = parsed[2:] if re.match(r"^:?-+", parsed[1][0]) else parsed[1:]
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    widths = [int(9360 / cols)] * cols
    if cols == 2:
        widths = [2650, 6710]
    elif cols == 3:
        widths = [1750, 3000, 4610]
    set_table_fixed_width(table, widths)
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "DDE8FF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        format_inline(p, text, color=NAVY, size=9.25, bold=True)
    for row_data in body:
        cells = table.add_row().cells
        for i in range(cols):
            text = row_data[i] if i < len(row_data) else ""
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[i], LIGHT)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            format_inline(p, text, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)
    add_cover(doc)

    in_code = False
    code_lines = []
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("# ") or line.startswith("> Last updated") or line.startswith("> Repository") or line.startswith("> Production") or line.startswith("> Local workspace"):
            index += 1
            continue
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                table = doc.add_table(rows=1, cols=1)
                set_table_fixed_width(table, [9360])
                cell = table.cell(0, 0)
                set_cell_shading(cell, "101A33")
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.6)
                run.font.color.rgb = WHITE
                in_code = False
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(lines[index])
                index += 1
            add_markdown_table(doc, rows)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            format_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            format_inline(p, line[2:])
        elif line.startswith("> "):
            add_callout(doc, line[2:], "Note")
        elif line:
            p = doc.add_paragraph()
            format_inline(p, line)
        index += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
